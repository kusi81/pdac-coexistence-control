"""manuscript.md -> manuscript.docx (그림 임베드·표·참조 포함).

블록 단위 파서: 하드 줄바꿈된 문단을 먼저 합치고(공백 줄=문단 경계), 헤딩/리스트/
표/인용/그림을 구조 처리. python-docx 조판. bioRxiv/리뷰용 단일 문서(그림 임베드).
"""
import os, re, sys
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(ROOT, "docs", "manuscript", "manuscript.md")
CLEAN = "--clean" in sys.argv   # 제출용: 내부 메모(draft note, 편집 체크리스트) 제외
OUT = os.path.join(ROOT, "docs", "manuscript",
                   "manuscript_clean.docx" if CLEAN else "manuscript.docx")
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`|\[.+?\]\(.+?\))")


def add_runs(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            p.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            p.add_run(tok[1:-1]).font.name = "Consolas"
        elif tok.startswith("[") and "](" in tok:
            r = p.add_run(tok[1:tok.index("](")])
            r.font.color.rgb = RGBColor(0x1A, 0x5A, 0x9A)
        else:
            p.add_run(tok)


def is_sep(l):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", l)) and "-" in l


def row(l):
    return [c.strip() for c in l.strip().strip("|").split("|")]


def blocks(lines):
    """(kind, payload) 블록 리스트로 그룹화."""
    out, para, i, n = [], [], 0, len(lines)
    def flush():
        if para:
            out.append(("p", " ".join(para))); para.clear()
    while i < n:
        s = lines[i].rstrip()
        st = s.strip()
        if st == "":
            flush(); i += 1
        elif s.startswith("#"):
            flush()
            lvl = len(s) - len(s.lstrip("#"))
            out.append(("h", (lvl, s.lstrip("#").strip()))); i += 1
        elif s.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            flush()
            hdr = row(s); i += 2; rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(row(lines[i])); i += 1
            out.append(("table", (hdr, rows)))
        elif re.match(r"^\d+\.\s", s):
            flush(); out.append(("ol", re.sub(r"^\d+\.\s", "", s))); i += 1
        elif s.startswith("- "):
            flush(); out.append(("ul", s[2:])); i += 1
        elif st == "---":
            flush(); i += 1
        elif s.startswith("> "):
            flush(); out.append(("quote", s[2:])); i += 1
        else:
            para.append(st); i += 1
    flush()
    return out


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.4; st.paragraph_format.space_after = Pt(6)

    for kind, pl in blocks(open(MD, encoding="utf-8").read().splitlines()):
        # 제출용 클린: 내부 메모 제외
        if CLEAN and kind == "h" and "Master editorial checklist" in pl[1]:
            break   # 이후(체크리스트) 전부 생략
        if CLEAN and kind == "p" and pl.lstrip("*").startswith("Draft manuscript"):
            continue
        if kind == "h":
            lvl, txt = pl
            add_runs(doc.add_heading("", level=min(lvl - 1, 4)), txt)
        elif kind == "table":
            hdr, rows = pl
            t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Light Grid Accent 1"
            for j, h in enumerate(hdr):
                c = t.rows[0].cells[j]; c.paragraphs[0].clear()
                add_runs(c.paragraphs[0], h)
                for r in c.paragraphs[0].runs:
                    r.bold = True
            for rr in rows:
                cs = t.add_row().cells
                for j, v in enumerate(rr[:len(hdr)]):
                    cs[j].paragraphs[0].clear(); add_runs(cs[j].paragraphs[0], v)
            doc.add_paragraph()
        elif kind == "ol":
            add_runs(doc.add_paragraph(style="List Number"), pl)
        elif kind == "ul":
            add_runs(doc.add_paragraph(style="List Bullet"), pl)
        elif kind == "quote":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
            add_runs(p, pl)
            for r in p.runs:
                r.italic = True
        else:  # p — 그림 캡션 특수 처리
            m = re.search(r"\((assets/[\w./-]+\.png)\)", pl)
            if m and pl.lstrip("*").startswith("Figure"):
                img = os.path.join(ROOT, m.group(1))
                if os.path.exists(img):
                    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pp.add_run().add_picture(img, width=Inches(6.0))
                cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(12)
                add_runs(cap, re.sub(r"\s*\(assets/[\w./-]+\.png\)\s*", "", pl))
                for r in cap.runs:
                    r.font.size = Pt(9.5)
            else:
                add_runs(doc.add_paragraph(), pl)

    doc.save(OUT)
    print("wrote", OUT, "| paras:", len(doc.paragraphs), "| tables:", len(doc.tables))


if __name__ == "__main__":
    main()
